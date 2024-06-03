# forum/utils.py

import requests
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO



def get_location_from_ip(ip_address):
    API_KEY = '69.166.46.175'
    url = f'http://api.ipstack.com/{ip_address}?access_key={API_KEY}'
    response = requests.get(url)
    data = response.json()
    try:
        city = data.get('city')
        if city:
            print(city)
            return city
        else:
            return None
    except Exception as e:
        print(f"Error: {e}")
        return None, None



def generate_resident_word(resident):
    document = Document()

    # Download logo image
    image_url = "https://cdn.discordapp.com/attachments/697362904130912259/1247232117557952542/image.png?ex=665f46ed&is=665df56d&hm=5ac6bc5a3c54e2eaeddac4e0a72541b4dac634cbede32d45307f9a91c6cd4dd5&"
    response = requests.get(image_url)
    image = BytesIO(response.content)

    # Add logo image to the document
    document.add_picture(image, width=Inches(2))

    # Corporate Header
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Park Rose Care Center"
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Title
    title = document.add_heading(level=1)
    run = title.add_run("Resident Inventory Form")
    run.font.size = Pt(24)
    run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Information
    document.add_heading('Resident Information', level=2)
    p = document.add_paragraph()
    p.add_run("ID: ").bold = True
    p.add_run(f"{resident.id}")
    p = document.add_paragraph()
    p.add_run("Name: ").bold = True
    p.add_run(f"{resident.name}")
    p = document.add_paragraph()
    p.add_run("Date of Admittance: ").bold = True
    p.add_run(f"{resident.date_of_admittance}")

    # Items
    document.add_heading('Items', level=2)
    items = resident.items.all()
    if items:
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item Name'
        hdr_cells[1].text = 'Signature'

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.name
            row_cells[1].text = ''

        # Add empty rows for additional sign-offs
        for _ in range(5):  # Add 5 additional rows for future items/signatures
            row_cells = table.add_row().cells
            row_cells[0].text = ''
            row_cells[1].text = ''
    else:
        document.add_paragraph("No items found.")

    # Comments
    document.add_heading('Comments', level=2)
    comments = resident.comments.all()
    if comments:
        for comment in comments:
            document.add_paragraph(comment.content, style='List Bullet')
    else:
        document.add_paragraph("No comments found.")

    # Signature lines without a table
    document.add_heading('Signatures', level=2)

    p = document.add_paragraph()
    p.add_run("Resident / Family Member: ").bold = True
    run = p.add_run(" " * 5 + "_" * 40)
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p = document.add_paragraph()
    p.add_run("Staff Signature: ").bold = True
    run = p.add_run(" " * 12 + "_" * 40)
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Footer
    footer = document.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Park Rose Care Center - Resident Inventory"
    footer_run = footer_paragraph.runs[0]
    footer_run.font.size = Pt(10)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Create a response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=resident_{resident.id}.docx'
    document.save(response)

    return response